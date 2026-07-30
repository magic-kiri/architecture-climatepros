#!/usr/bin/env python3
"""Build a client-ready, professionally designed .docx from
stream1-unified-architecture-doc.html.

Design system: corporate navy + warm copper accent (tokens below).
The html export flattens several structures the live document has (stat strips,
numbered stack lists, flow chains, context chips) into single <p> blocks with
soft line breaks. This script re-detects and re-typesets them.

Run:  python3 scripts/build_docx.py
Out:  stream1-unified-architecture-doc.docx
Two-pass: pass 1 builds with placeholder TOC page numbers, renders a PDF to
learn the real page of every heading, pass 2 writes the final file.
"""
import os
import re
import pathlib
import shutil
import subprocess
import sys
from collections import namedtuple

from bs4 import BeautifulSoup, NavigableString
import docx
from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, Inches, RGBColor, Emu

ROOT = pathlib.Path(__file__).resolve().parent.parent
SRC = ROOT / 'stream1-unified-architecture-doc.html'
OUT = ROOT / 'stream1-unified-architecture-doc.docx'
WORK = pathlib.Path('/tmp/docxbuild')
FIG_DIR = ROOT / 'assets' / 'figures'      # produced by scripts/capture_figures.py

# how much room a figure can have, in inches, for each placement option
# (max width, max image height) — heights leave room for the frame's label,
# title and caption so the whole plate always fits on one page
FIT_INLINE = (6.5, 7.1)           # in the flow of the section, framed
FIT_LANDSCAPE_PAGE = (9.4, 5.15)  # alone on a rotated page

# --------------------------------------------------------------- design tokens
# Palette, type and brand furniture lifted from sample-unified.docx
BRAND = ROOT / 'assets' / 'brand'
FONT_DIR = ROOT / 'assets' / 'fonts'

NAVY = '1F3864'          # headings, table headers, section numbers
NAVY_MID = '2E4C7E'      # subheadings, inline code, secondary rules
VIOLET = '6A5BF5'        # accent, from the header banner's gradient
VIOLET_DEEP = '5348D0'
TEAL = '2E4C7E'          # kept as an alias so use-case C stays in-palette
COPPER = VIOLET          # legacy name -> brand accent
COPPER_LT = 'DEDBFA'
RED = 'A3312A'
INK = '2A3342'           # body text
MUTED = '5B6472'         # captions, meta, muted labels
RULE = 'CFD9E8'
RULE_LT = 'E3EBF5'
SOFT = 'EEF3F9'          # zebra rows, tints
CODE_BG = 'F2F6FB'
FLOW_BG = 'E9F0F9'
CALLOUT_BG = 'F1F3FE'
FIG_BG = 'F8FAFD'
WHITE = 'FFFFFF'
NAVY_SOFT = 'E3EBF5'

BADGE = {'A': NAVY, 'B': NAVY_MID, 'C': VIOLET}

SERIF = 'Poppins'        # display face (embedded in the output document)
SANS = 'Calibri'
MONO = 'Consolas'

BODY_PT = 10.5
TABLE_PT = 9.0
TABLE_PT_TIGHT = 8.0
CONTENT_W = Inches(6.5)
DOC_TITLE = 'ClimatePros · Stream 1 — Unified Architecture'

KEEP_GLYPHS = set('✓✔✗✘×•')
EMOJI = re.compile('[\U0001F300-\U0001FAFF\U00002600-\U000027BF\U0001F000-\U0001F0FF'
                   '\U00002B00-\U00002BFF️]')


def strip_emoji(text):
    return EMOJI.sub(lambda m: m.group(0) if m.group(0) in KEEP_GLYPHS else '', text)
ARROWS = '→↔⇄↻↓⇒'


def clean(text):
    text = strip_emoji(text)
    text = text.replace(' ', ' ').replace(' ', ' ')
    text = re.sub(r'[ \t]{2,}', ' ', text)
    return text


# ------------------------------------------------------------------- xml helpers
def _el(tag, **attrs):
    e = OxmlElement(tag)
    for k, v in attrs.items():
        e.set(qn('w:' + k), str(v))
    return e


# ---- schema-ordered insertion -------------------------------------------------
# CT_* types in WordprocessingML are xsd:sequence, so a property element has to
# be placed in canonical order or Word flags the file as needing repair.
ORDER = {
    'pPr': ['pStyle', 'keepNext', 'keepLines', 'pageBreakBefore', 'framePr',
            'widowControl', 'numPr', 'suppressLineNumbers', 'pBdr', 'shd', 'tabs',
            'suppressAutoHyphens', 'kinsoku', 'wordWrap', 'overflowPunct',
            'topLinePunct', 'autoSpaceDE', 'autoSpaceDN', 'bidi', 'adjustRightInd',
            'snapToGrid', 'spacing', 'ind', 'contextualSpacing', 'mirrorIndents',
            'suppressOverlap', 'jc', 'textDirection', 'textAlignment',
            'textboxTightWrap', 'outlineLvl', 'divId', 'cnfStyle', 'rPr', 'sectPr',
            'pPrChange'],
    'rPr': ['rStyle', 'rFonts', 'b', 'bCs', 'i', 'iCs', 'caps', 'smallCaps',
            'strike', 'dstrike', 'outline', 'shadow', 'emboss', 'imprint', 'noProof',
            'snapToGrid', 'vanish', 'webHidden', 'color', 'spacing', 'w', 'kern',
            'position', 'sz', 'szCs', 'highlight', 'u', 'effect', 'bdr', 'shd',
            'fitText', 'vertAlign', 'rtl', 'cs', 'em', 'lang', 'eastAsianLayout',
            'specVanish', 'oMath'],
    'tcPr': ['cnfStyle', 'tcW', 'gridSpan', 'hMerge', 'vMerge', 'tcBorders', 'shd',
             'noWrap', 'tcMar', 'textDirection', 'tcFitText', 'vAlign', 'hideMark'],
    'tblPr': ['tblStyle', 'tblpPr', 'tblOverlap', 'bidiVisual', 'tblStyleRowBandSize',
              'tblStyleColBandSize', 'tblW', 'jc', 'tblCellSpacing', 'tblInd',
              'tblBorders', 'shd', 'tblLayout', 'tblCellMar', 'tblLook', 'tblCaption',
              'tblDescription'],
    'trPr': ['cnfStyle', 'divId', 'gridBefore', 'gridAfter', 'wBefore', 'wAfter',
             'cantSplit', 'trHeight', 'tblHeader', 'tblCellSpacing', 'jc', 'hidden'],
    'sectPr': ['footnotePr', 'endnotePr', 'type', 'pgSz', 'pgMar', 'paperSrc',
               'pgBorders', 'lnNumType', 'pgNumType', 'cols', 'formProt', 'vAlign',
               'noEndnote', 'titlePg', 'textDirection', 'bidi', 'rtlGutter',
               'docGrid', 'printerSettings', 'sectPrChange'],
    'settings': ['writeProtection', 'view', 'zoom', 'removePersonalInformation',
                 'removeDateAndTime', 'doNotDisplayPageBoundaries',
                 'displayBackgroundShape', 'printPostScriptOverText',
                 'printFractionalCharacterWidth', 'printFormsData', 'embedTrueTypeFonts',
                 'embedSystemFonts', 'saveSubsetFonts', 'saveFormsData',
                 'mirrorMargins', 'alignBordersAndEdges', 'bordersDoNotSurroundHeader',
                 'bordersDoNotSurroundFooter', 'gutterAtTop', 'hideSpellingErrors',
                 'hideGrammaticalErrors', 'activeWritingStyle', 'proofState',
                 'formsDesign', 'attachedTemplate', 'linkStyles',
                 'stylePaneFormatFilter', 'stylePaneSortMethod', 'documentType',
                 'mailMerge', 'revisionView', 'trackChanges', 'documentProtection',
                 'autoFormatOverride', 'styleLockTheme', 'styleLockQFSet',
                 'defaultTabStop', 'autoHyphenation', 'consecutiveHyphenLimit',
                 'hyphenationZone', 'doNotHyphenateCaps', 'showEnvelope',
                 'summaryLength', 'clickAndTypeStyle', 'defaultTableStyle',
                 'evenAndOddHeaders', 'bookFoldRevPrinting', 'bookFoldPrinting',
                 'bookFoldPrintingSheets', 'drawingGridHorizontalSpacing',
                 'drawingGridVerticalSpacing', 'displayHorizontalDrawingGridEvery',
                 'displayVerticalDrawingGridEvery', 'doNotUseMarginsForDrawingGridOrigin',
                 'drawingGridHorizontalOrigin', 'drawingGridVerticalOrigin',
                 'doNotShadeFormData', 'noPunctuationKerning',
                 'characterSpacingControl', 'printTwoOnOne', 'strictFirstAndLastChars',
                 'noLineBreaksAfter', 'noLineBreaksBefore', 'savePreviewPicture',
                 'doNotValidateAgainstSchema', 'saveInvalidXml', 'ignoreMixedContent',
                 'alwaysShowPlaceholderText', 'doNotDemarcateInvalidXml',
                 'saveXmlDataOnly', 'useXSLTWhenSaving', 'saveThroughXslt',
                 'showXMLTags', 'alwaysMergeEmptyNamespace', 'updateFields',
                 'hdrShapeDefaults', 'footnotePr', 'endnotePr', 'compat', 'docVars',
                 'rsids'],
}


def put(parent, element):
    """Insert element into parent at its schema position (replacing any existing)."""
    tag = element.tag.split('}')[1]
    ptag = parent.tag.split('}')[1]
    order = ORDER.get(ptag)
    old = parent.find(qn('w:' + tag))
    if old is not None:
        parent.remove(old)
    if not order or tag not in order:
        parent.append(element)
        return element
    rank = order.index(tag)
    for child in parent:
        ctag = child.tag.split('}')[1]
        if ctag in order and order.index(ctag) > rank:
            child.addprevious(element)
            return element
    parent.append(element)
    return element


def para_shade(p, fill):
    put(p._p.get_or_add_pPr(), _el('w:shd', val='clear', color='auto', fill=fill))


def cell_shade(cell, fill):
    put(cell._tc.get_or_add_tcPr(), _el('w:shd', val='clear', color='auto', fill=fill))


def para_borders(p, **edges):
    pPr = p._p.get_or_add_pPr()
    bdr = OxmlElement('w:pBdr')
    for name in ('top', 'left', 'bottom', 'right'):
        if name in edges and edges[name]:
            style, sz, color = edges[name]
            bdr.append(_el('w:' + name, val=style, sz=sz, space=6, color=color))
    put(pPr, bdr)


def cell_borders(cell, **edges):
    tcPr = cell._tc.get_or_add_tcPr()
    bdr = OxmlElement('w:tcBorders')
    for name in ('top', 'left', 'bottom', 'right'):
        if name in edges:
            spec = edges[name]
            if spec is None:
                bdr.append(_el('w:' + name, val='none', sz=0, space=0, color='auto'))
            else:
                style, sz, color = spec
                bdr.append(_el('w:' + name, val=style, sz=sz, space=0, color=color))
    put(tcPr, bdr)


def cell_margins(cell, top=60, left=110, bottom=60, right=110):
    tcPr = cell._tc.get_or_add_tcPr()
    mar = OxmlElement('w:tcMar')
    for name, val in (('top', top), ('left', left), ('bottom', bottom), ('right', right)):
        mar.append(_el('w:' + name, w=val, type='dxa'))
    put(tcPr, mar)


def table_borders(table, inside_h=('single', 4, RULE), inside_v=('single', 4, RULE_LT),
                  outer=('single', 4, RULE)):
    tblPr = table._tbl.tblPr
    b = OxmlElement('w:tblBorders')
    spec = {'top': outer, 'left': outer, 'bottom': outer, 'right': outer,
            'insideH': inside_h, 'insideV': inside_v}
    for name in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        s = spec[name]
        if s is None:
            b.append(_el('w:' + name, val='none', sz=0, space=0, color='auto'))
        else:
            style, sz, color = s
            b.append(_el('w:' + name, val=style, sz=sz, space=0, color=color))
    put(tblPr, b)


def repeat_header(row):
    put(row._tr.get_or_add_trPr(), _el('w:tblHeader', val='true'))


def cant_split(row):
    put(row._tr.get_or_add_trPr(), _el('w:cantSplit', val='true'))


def row_height(row, inches, rule='atLeast'):
    put(row._tr.get_or_add_trPr(),
        _el('w:trHeight', val=int(inches * 1440), hRule=rule))


def letterspace(run, twentieths=40):
    put(run._element.get_or_add_rPr(), _el('w:spacing', val=twentieths))


def run_shade(run, fill):
    put(run._element.get_or_add_rPr(),
        _el('w:shd', val='clear', color='auto', fill=fill))


def add_field(paragraph, instr, placeholder='1'):
    runs = []
    r = paragraph.add_run(); r._r.append(_el('w:fldChar', fldCharType='begin')); runs.append(r)
    r = paragraph.add_run()
    it = OxmlElement('w:instrText'); it.set(qn('xml:space'), 'preserve'); it.text = instr
    r._r.append(it); runs.append(r)
    r = paragraph.add_run(); r._r.append(_el('w:fldChar', fldCharType='separate')); runs.append(r)
    r = paragraph.add_run(placeholder); runs.append(r)
    r = paragraph.add_run(); r._r.append(_el('w:fldChar', fldCharType='end')); runs.append(r)
    return runs


def set_tbl_w(table, width=CONTENT_W):
    tblPr = table._tbl.tblPr
    put(tblPr, _el('w:tblW', w=int(width.twips), type='dxa'))
    put(tblPr, _el('w:tblInd', w=0, type='dxa'))
    put(tblPr, _el('w:tblLayout', type='fixed'))


# ---------------------------------------------------------------------- styles
def build_styles(doc):
    st = doc.styles
    normal = st['Normal']
    normal.font.name = SANS
    normal.font.size = Pt(BODY_PT)
    normal.font.color.rgb = RGBColor.from_string(INK)
    pf = normal.paragraph_format
    pf.space_after = Pt(6.5)
    pf.line_spacing = 1.25
    pf.widow_control = True

    def heading(name, size, color, font=SERIF, bold=True, before=0, after=6, italic=False):
        s = st[name]
        s.font.name = font
        s.font.size = Pt(size)
        s.font.bold = bold
        s.font.italic = italic
        s.font.color.rgb = RGBColor.from_string(color)
        s.paragraph_format.space_before = Pt(before)
        s.paragraph_format.space_after = Pt(after)
        s.paragraph_format.line_spacing = 1.1
        s.paragraph_format.keep_with_next = True
        s.paragraph_format.keep_together = True

    heading('Heading 1', 21, NAVY, after=13)
    heading('Heading 2', 13.5, NAVY_MID, before=18, after=5)
    heading('Heading 3', 10.5, NAVY_MID, font=SANS, before=13, after=2)
    heading('Heading 4', 10, INK, font=SANS, before=10, after=1)
    heading('Heading 5', 9.5, MUTED, font=SANS, before=9, after=1, italic=True)
    st['Heading 1'].paragraph_format.page_break_before = True

    def new(name, base='Normal'):
        try:
            s = st.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
        except ValueError:
            s = st[name]
        s.base_style = st[base]
        s.quick_style = False
        return s

    s = new('CoverEyebrow'); f = s.font
    f.name, f.size, f.bold, f.all_caps = SANS, Pt(8.5), True, True
    f.color.rgb = RGBColor.from_string(COPPER)
    s.paragraph_format.space_after = Pt(4)

    s = new('CoverTitle'); f = s.font
    f.name, f.size, f.bold = SERIF, Pt(31), True
    f.color.rgb = RGBColor.from_string(WHITE)
    s.paragraph_format.space_after = Pt(2)
    s.paragraph_format.line_spacing = 1.02

    s = new('CoverSub'); f = s.font
    f.name, f.size = SANS, Pt(11)
    f.color.rgb = RGBColor.from_string('C7D2E4')
    s.paragraph_format.space_after = Pt(3)

    s = new('MetaLabel'); f = s.font
    f.name, f.size, f.bold, f.all_caps = SANS, Pt(7.5), True, True
    f.color.rgb = RGBColor.from_string(COPPER)
    s.paragraph_format.space_after = Pt(1)
    s.paragraph_format.line_spacing = 1.0

    s = new('MetaValue'); f = s.font
    f.name, f.size = SANS, Pt(10)
    f.color.rgb = RGBColor.from_string(NAVY)
    s.paragraph_format.space_after = Pt(11)
    s.paragraph_format.line_spacing = 1.15

    s = new('BlockTitle'); f = s.font
    f.name, f.size, f.bold = SERIF, Pt(15), True
    f.color.rgb = RGBColor.from_string(NAVY)
    s.paragraph_format.space_before = Pt(16)
    s.paragraph_format.space_after = Pt(9)
    s.paragraph_format.keep_with_next = True

    s = new('Bullet')
    s.paragraph_format.left_indent = Inches(0.26)
    s.paragraph_format.first_line_indent = Inches(-0.17)
    s.paragraph_format.space_after = Pt(3.5)

    s = new('Bullet2')
    s.paragraph_format.left_indent = Inches(0.55)
    s.paragraph_format.first_line_indent = Inches(-0.17)
    s.paragraph_format.space_after = Pt(3)

    s = new('Code'); f = s.font
    f.name, f.size = MONO, Pt(8.5)
    f.color.rgb = RGBColor.from_string('1F2A37')
    s.paragraph_format.space_after = Pt(0)
    s.paragraph_format.space_before = Pt(0)
    s.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

    s = new('Flow'); f = s.font
    f.name, f.size, f.bold = MONO, Pt(8.5), True
    f.color.rgb = RGBColor.from_string(NAVY_MID)
    s.paragraph_format.space_before = Pt(4)
    s.paragraph_format.space_after = Pt(7)
    s.paragraph_format.left_indent = Inches(0.09)
    s.paragraph_format.line_spacing = 1.1

    s = new('Chip'); f = s.font
    f.name, f.size, f.bold, f.small_caps = SANS, Pt(8), True, True
    f.color.rgb = RGBColor.from_string(MUTED)
    s.paragraph_format.space_before = Pt(9)
    s.paragraph_format.space_after = Pt(2)
    s.paragraph_format.keep_with_next = True

    s = new('StepHead'); f = s.font
    f.name, f.size, f.bold = SANS, Pt(10), True
    f.color.rgb = RGBColor.from_string(NAVY_MID)
    s.paragraph_format.space_before = Pt(9)
    s.paragraph_format.space_after = Pt(3)
    s.paragraph_format.keep_with_next = True

    s = new('CellText'); f = s.font
    f.size = Pt(TABLE_PT)
    s.paragraph_format.space_after = Pt(0)
    s.paragraph_format.space_before = Pt(0)
    s.paragraph_format.line_spacing = 1.16

    s = new('CellHead'); f = s.font
    f.size, f.bold = Pt(TABLE_PT), True
    f.color.rgb = RGBColor.from_string(WHITE)
    s.paragraph_format.space_after = Pt(0)
    s.paragraph_format.space_before = Pt(0)
    s.paragraph_format.line_spacing = 1.12

    s = new('MetricValue'); f = s.font
    f.name, f.size, f.bold = SERIF, Pt(15), True
    f.color.rgb = RGBColor.from_string(NAVY)
    s.paragraph_format.space_after = Pt(1)
    s.paragraph_format.space_before = Pt(0)
    s.paragraph_format.line_spacing = 1.0

    s = new('MetricLabel'); f = s.font
    f.name, f.size, f.bold, f.all_caps = SANS, Pt(7.5), True, True
    f.color.rgb = RGBColor.from_string(COPPER)
    s.paragraph_format.space_after = Pt(0)
    s.paragraph_format.space_before = Pt(0)
    s.paragraph_format.line_spacing = 1.0

    s = new('MetricNote'); f = s.font
    f.name, f.size = SANS, Pt(8.5)
    f.color.rgb = RGBColor.from_string(MUTED)
    s.paragraph_format.space_after = Pt(0)
    s.paragraph_format.space_before = Pt(0)
    s.paragraph_format.line_spacing = 1.14

    s = new('Caption'); f = s.font
    f.size = Pt(8.5)
    f.color.rgb = RGBColor.from_string(MUTED)
    s.paragraph_format.space_after = Pt(0)
    s.paragraph_format.line_spacing = 1.2

    s = new('FigLabel'); f = s.font
    f.name, f.size, f.bold, f.all_caps = SANS, Pt(8), True, True
    f.color.rgb = RGBColor.from_string(COPPER)
    s.paragraph_format.space_after = Pt(2)

    s = new('FigTitle'); f = s.font
    f.name, f.size, f.bold = SANS, Pt(10.5), True
    f.color.rgb = RGBColor.from_string(NAVY)
    s.paragraph_format.space_after = Pt(4)

    s = new('FigDrop'); f = s.font
    f.size, f.italic = Pt(9), True
    f.color.rgb = RGBColor.from_string('9AA7B7')
    s.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    s.paragraph_format.space_before = Pt(30)
    s.paragraph_format.space_after = Pt(30)

    s = new('Callout'); f = s.font
    f.size = Pt(10)
    s.paragraph_format.space_after = Pt(3)
    s.paragraph_format.line_spacing = 1.2

    s = new('TocH1'); f = s.font
    f.name, f.size, f.bold = SANS, Pt(10.5), True
    f.color.rgb = RGBColor.from_string(NAVY)
    s.paragraph_format.space_before = Pt(11)
    s.paragraph_format.space_after = Pt(2)
    s.paragraph_format.line_spacing = 1.0

    s = new('TocH2'); f = s.font
    f.name, f.size = SANS, Pt(9.5)
    f.color.rgb = RGBColor.from_string(INK)
    s.paragraph_format.space_after = Pt(1.5)
    s.paragraph_format.left_indent = Inches(0.3)
    s.paragraph_format.line_spacing = 1.0

    s = new('HeadFoot'); f = s.font
    f.name, f.size = SANS, Pt(8)
    f.color.rgb = RGBColor.from_string(MUTED)
    s.paragraph_format.space_after = Pt(0)
    s.paragraph_format.space_before = Pt(0)
    s.paragraph_format.line_spacing = 1.0


# ------------------------------------------------------- inline tokens & output
Tok = namedtuple('Tok', 'text bold italic code href')
NO_SPACE_BEFORE = set(',.;:!?)]}%’\'"')


def tokenize(node):
    """-> list of lines; each line is a list of Tok. Soft \n becomes a new line."""
    lines = [[]]

    def walk(n, bold, italic, code, href):
        for child in n.children:
            if isinstance(child, NavigableString):
                raw = clean(str(child))
                if not raw:
                    continue
                parts = raw.split('\n')
                for i, part in enumerate(parts):
                    if i:
                        lines.append([])
                    if part == '':
                        continue
                    if i:
                        part = part.lstrip()
                    lines[-1].append(Tok(part, bold, italic, code, href))
                continue
            name = child.name
            if name in ('strong', 'b'):
                walk(child, True, italic, code, href)
            elif name in ('em', 'i'):
                walk(child, bold, True, code, href)
            elif name == 'code':
                walk(child, bold, italic, True, href)
            elif name == 'br':
                lines.append([])
            elif name == 'a':
                h = child.get('href', '')
                walk(child, bold, italic, code, h if h.startswith('http') else href)
            else:
                walk(child, bold, italic, code, href)

    walk(node, False, False, False, None)
    out = []
    for ln in lines:
        toks = [t for t in ln if t.text.strip() or t.text == ' ']
        if toks:
            out.append(toks)
    return out


def line_text(toks):
    s = ''.join(t.text for t in toks)
    return tidy(s)


def tidy(s):
    s = re.sub(r'\s+([,.;:!?%])(?![A-Za-z0-9])', r'\1', s)
    s = re.sub(r'\(\s+', '(', s)
    s = re.sub(r'\s+\)', ')', s)
    s = re.sub(r'[ ]{2,}', ' ', s)
    return s.strip()


def emit(p, toks, size=None, color=None, code_size=None, bold_all=False):
    prev = None
    for t in toks:
        txt = re.sub(r'(?<=[\d)])\s+%', '%', t.text)
        txt = re.sub(r'\s+([,.;:!?])(?![A-Za-z0-9])', r'\1', txt)
        txt = re.sub(r'\(\s+', '(', txt)
        txt = re.sub(r'\s+\)', ')', txt)
        # a token starting with punctuation usually means the space before it is
        # a stray from the export — but not when the punctuation opens a word
        # (".NET", ".p8"), where the space is real
        opens_word = re.match(r'^[.,;:!?)\]}%]\w', txt) is not None
        pct_after_word = txt[:1] == '%' and not re.search(r'\d\s*$', prev.text if prev else '')
        if (prev is not None and txt[:1] in NO_SPACE_BEFORE and not opens_word
                and not pct_after_word and prev.text.endswith(' ')):
            # remove the stray space the html export left before punctuation
            for r in reversed(p.runs):
                if r.text.endswith(' '):
                    r.text = r.text[:-1]
                    break
        if prev is None and not p.runs:
            txt = txt.lstrip()
        if not txt:
            continue
        if t.href:
            add_hyperlink(p, txt.strip(), t.href, size)
            prev = t
            continue
        r = p.add_run(txt)
        r.bold = t.bold or bold_all
        r.italic = t.italic
        if t.code:
            r.font.name = MONO
            r.font.size = Pt(code_size or (size - 1.2 if size else 9))
            r.font.color.rgb = RGBColor.from_string(NAVY_MID)
            run_shade(r, 'E9EDF3')
        else:
            if size:
                r.font.size = Pt(size)
            if color:
                r.font.color.rgb = RGBColor.from_string(color)
        prev = t
    # trailing space cleanup
    for r in reversed(p.runs):
        if r.text.endswith(' '):
            r.text = r.text.rstrip()
        break
    return p


def add_hyperlink(p, text, url, size=None):
    r_id = p.part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK,
                            is_external=True)
    link = OxmlElement('w:hyperlink')
    link.set(qn('r:id'), r_id)
    r = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    c = OxmlElement('w:color'); c.set(qn('w:val'), NAVY_MID); rPr.append(c)
    u = OxmlElement('w:u'); u.set(qn('w:val'), 'single'); rPr.append(u)
    if size:
        sz = OxmlElement('w:sz'); sz.set(qn('w:val'), str(int(size * 2))); rPr.append(sz)
    r.append(rPr)
    t = OxmlElement('w:t'); t.text = text; r.append(t)
    link.append(r)
    p._p.append(link)


def text_of(node):
    return tidy(clean(node.get_text(' ', strip=True)))


# -------------------------------------------------------------------- components
def one_cell(doc, fill, borders, margins=(90, 140, 90, 140)):
    t = doc.add_table(rows=1, cols=1)
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    t.autofit = False
    set_tbl_w(t)
    t.columns[0].width = CONTENT_W
    cell = t.cell(0, 0)
    cell.width = CONTENT_W
    cell_shade(cell, fill)
    cell_borders(cell, **borders)
    cell_margins(cell, *margins)
    return t, cell


def vgap(doc, points):
    """A reliable vertical gap: an empty run's font size is ignored by some
    renderers, so the height is pinned with exact line spacing instead."""
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.line_spacing = Pt(points)
    r = p.add_run(' ')
    r.font.size = Pt(1)
    return p


def spacer(doc, pts=8):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.line_spacing = 1.0
    p.add_run('').font.size = Pt(pts)
    return p


def render_code(doc, node):
    lines = [strip_emoji(ln) for ln in
             node.get_text().replace(' ', ' ').rstrip('\n').split('\n')]
    t, cell = one_cell(doc, CODE_BG, {
        'left': ('single', 18, NAVY_MID), 'top': ('single', 4, RULE),
        'bottom': ('single', 4, RULE), 'right': ('single', 4, RULE)},
        margins=(110, 150, 110, 130))
    for i, ln in enumerate(lines):
        p = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
        p.style = doc.styles['Code']
        p.add_run(ln)
    if len(lines) <= 20:            # taller blocks must be allowed to break
        cant_split(t.rows[0])
    spacer(doc, 9)


def render_callout(doc, node):
    t, cell = one_cell(doc, CALLOUT_BG, {
        'left': ('single', 18, COPPER), 'top': ('single', 4, COPPER_LT),
        'bottom': ('single', 4, COPPER_LT), 'right': ('single', 4, COPPER_LT)},
        margins=(110, 150, 110, 140))
    used = [False]

    def para(style='Callout'):
        p = cell.paragraphs[0] if not used[0] else cell.add_paragraph()
        used[0] = True
        p.style = doc.styles[style]
        return p

    for child in node.children:
        if child.name == 'p':
            for ln in tokenize(child):
                emit(para(), ln, size=10)
        elif child.name in ('ul', 'ol'):
            for li in child.find_all('li', recursive=False):
                p = para()
                p.paragraph_format.left_indent = Inches(0.2)
                p.paragraph_format.first_line_indent = Inches(-0.14)
                r = p.add_run('— ')
                r.font.color.rgb = RGBColor.from_string(COPPER)
                for ln in tokenize(li):
                    emit(p, ln, size=10)
                    break
        elif child.name == 'pre':
            p = para('Code')
            p.add_run(child.get_text().strip())
    cant_split(t.rows[0])
    spacer(doc, 9)


FIG_RE = re.compile(r'^Figure\s+(\d+)\s*[—–-]\s*(.*)$')


def png_size(path):
    data = path.read_bytes()[16:24]
    return int.from_bytes(data[:4], 'big'), int.from_bytes(data[4:], 'big')


def fit(aspect, box):
    """Largest (w, h) with this aspect that fits box=(max_w, max_h), in inches."""
    max_w, max_h = box
    w = min(max_w, max_h / aspect)
    return w, w * aspect


def figure_plan(num):
    """Where to put figure `num` and how big — whichever placement renders it
    widest wins, but a dedicated page has to earn its page break."""
    img = FIG_DIR / f'fig{num}.png'
    if not img.exists():
        return None
    px_w, px_h = png_size(img)
    aspect = px_h / px_w
    inline = fit(aspect, FIT_INLINE)
    land = fit(aspect, FIT_LANDSCAPE_PAGE)
    # a rotated page has to make the diagram meaningfully bigger to be worth it
    best = ('landscape', land) if land[0] > inline[0] * 1.15 else ('inline', inline)
    return {'path': img, 'placement': best[0], 'w': best[1][0], 'h': best[1][1],
            'px': (px_w, px_h), 'aspect': aspect}


def figure_frame(doc, width, plan, num, title, caption_node, styles):
    """The framed plate: label, title, image (or drop cue), caption."""
    solid = plan is not None
    edge = ('single', 4, RULE) if solid else ('dashed', 6, 'BCC7D6')
    t = doc.add_table(rows=1, cols=1)
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    t.autofit = False
    set_tbl_w(t, width)
    t.columns[0].width = width
    cell = t.cell(0, 0)
    cell.width = width
    cell_shade(cell, WHITE if solid else FIG_BG)
    cell_borders(cell, left=edge, top=edge, bottom=edge, right=edge)
    cell_margins(cell, 130, 150, 130, 150)

    p = cell.paragraphs[0]
    p.style = styles['FigLabel']
    letterspace(p.add_run(f'Figure {num}'), 32)
    p2 = cell.add_paragraph(style=styles['FigTitle'])
    p2.add_run(title)

    if solid:
        pic = cell.add_paragraph()
        pic.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pic.paragraph_format.space_before = Pt(4)
        pic.paragraph_format.space_after = Pt(7)
        pic.paragraph_format.line_spacing = 1.0
        pic.add_run().add_picture(str(plan['path']), width=Inches(plan['w']))
    else:
        drop = cell.add_paragraph(style=styles['FigDrop'])
        letterspace(drop.add_run('[  paste the diagram screenshot here  ]'), 20)

    if caption_node is not None:
        for em in caption_node.find_all('em'):
            if 'paste screenshot' in em.get_text(' ', strip=True).lower():
                em.extract()
        for i, ln in enumerate(tokenize(caption_node)):
            cap = cell.add_paragraph(style=styles['Caption'])
            emit(cap, ln, size=8.5, color=MUTED, code_size=8)
    return t


MD_FIG_SRC = re.compile(r'^fig(\d+)$', re.I)


def render_md_figure(doc, img):
    """![Title](figN "caption") -> the house figure plate."""
    m = MD_FIG_SRC.match((img.get('src') or '').strip())
    if not m:
        return False
    num = m.group(1)
    plan = figure_plan(num)
    title = (img.get('alt') or '').strip()
    caption = (img.get('title') or '').strip()
    cap_node = BeautifulSoup(f'<p>{caption}</p>', 'html.parser').p if caption else None
    if plan and plan['placement'] == 'landscape':
        new_page_section(doc, landscape=True)
        t = figure_frame(doc, Emu(int(Inches(9.6).emu)), plan, num, title, cap_node,
                         doc.styles)
        cant_split(t.rows[0])
        new_page_section(doc, landscape=False)
    else:
        t = figure_frame(doc, CONTENT_W, plan, num, title, cap_node, doc.styles)
        cant_split(t.rows[0])
        if not plan or plan['h'] < 6.0:
            spacer(doc, 10)
    return True


def render_figure_block(doc, node):
    paras = node.find_all('p', recursive=False)
    groups = []
    for p in paras:
        m = FIG_RE.match(text_of(p)) if p.find('strong') else None
        if m:
            groups.append({'num': m.group(1), 'title': m.group(2), 'caption': None})
        elif groups:
            groups[-1]['caption'] = p

    for g in groups:
        plan = figure_plan(g['num'])
        placement = plan['placement'] if plan else 'inline'
        if placement == 'inline':
            t = figure_frame(doc, CONTENT_W, plan, g['num'], g['title'],
                             g['caption'], doc.styles)
            # keep the plate whole: if it no longer fits here Word moves all of
            # it to the next page rather than orphaning the caption
            cant_split(t.rows[0])
            # a near-full-page plate must not push a trailing spacer onto the
            # next page, which would leave an empty sheet behind
            if not plan or plan['h'] < 6.0:
                spacer(doc, 10)
        else:
            # rotated plate on its own sheet, then back to portrait
            sec, width = new_page_section(doc, landscape=True)
            # centre the plate on the sheet: pad the top by half the slack
            cap_len = len(text_of(g['caption'])) if g['caption'] is not None else 0
            plate_h = (plan['h'] + 0.7
                       + 0.13 * max(1, -(-cap_len // int(plan['w'] * 17))))
            slack = max(0.0, (7.15 - plate_h) / 2)
            if slack > 0.6:
                spacer(doc, min(slack, 2.6) * 72)
            t = figure_frame(doc, width, plan, g['num'], g['title'], g['caption'],
                             doc.styles)
            cant_split(t.rows[0])
            new_page_section(doc, landscape=False)


def col_widths(rows, total=CONTENT_W):
    ncols = max(len(r) for r in rows)
    mass = [1.0] * ncols
    for r in rows:
        for i, cell in enumerate(r):
            mass[i] = max(mass[i], min(len(cell), 200) ** 0.6)
    tot = sum(mass)
    raw = [m / tot for m in mass]
    floor = 0.08 if ncols > 4 else 0.11
    raw = [max(x, floor) for x in raw]
    s = sum(raw)
    raw = [x / s for x in raw]
    widths = [Emu(int(total.emu * x)) for x in raw]
    widths[-1] = Emu(widths[-1].emu + (total.emu - sum(w.emu for w in widths)))
    return widths


NUMERIC_CELL = re.compile(r'^[~≈<>≤≥$€£+\-–—()\d.,%/×x\s]*$')


def numeric_columns(text_rows):
    """Money and count columns read far better right-aligned."""
    if len(text_rows) < 2:
        return set()
    ncols = max(len(r) for r in text_rows)
    right = set()
    for i in range(ncols):
        cells = [r[i] for r in text_rows[1:] if i < len(r) and r[i].strip()]
        if not cells:
            continue
        numish = [c for c in cells if NUMERIC_CELL.match(c) and any(ch.isdigit() for ch in c)]
        if len(numish) >= max(2, int(0.6 * len(cells))):
            right.add(i)
    return right


def is_summary_row(cells_html):
    """A row whose every filled cell is entirely bold — a totals line."""
    filled = [c for c in cells_html if text_of(c)]
    if not filled:
        return False
    for cell in filled:
        text = text_of(cell)
        bold = ' '.join(b.get_text(' ', strip=True) for b in cell.find_all('strong'))
        if not bold or len(clean(bold)) < len(text) - 2:
            return False
    return True


def render_summary_band(doc, head_texts, total_cells, widths):
    """Headline the totals row on the same column grid as the table below it:
    lead label, then each total set large over its column heading."""
    values = [text_of(c) for c in total_cells]
    if len(values) != len(widths) or len(values) < 3:
        return
    t = doc.add_table(rows=1, cols=len(widths))
    t.autofit = False
    set_tbl_w(t)
    for i, w in enumerate(widths):
        t.columns[i].width = w
    table_borders(t, inside_h=None, inside_v=None, outer=None)

    # one type size for the whole band, shrunk until the widest total fits
    size = 16.0
    for i in range(1, len(values)):
        avail = widths[i].inches - 0.10        # cell padding below
        while size > 9 and len(values[i]) * 0.66 * size / 72 > avail:
            size -= 0.5

    for i in range(len(widths)):
        cell = t.cell(0, i)
        cell.width = widths[i]
        cell_margins(cell, 80, 40, 70, 40)
        cell_borders(cell, top=('single', 10, NAVY), bottom=('single', 4, RULE),
                     left=None, right=None)
        p = cell.paragraphs[0]
        if i == 0:
            p.style = doc.styles['MetricLabel']
            p.paragraph_format.space_before = Pt(7)
            letterspace(p.add_run(re.sub(r'\s*\(est\.\)$', '', values[0])), 28)
            continue
        p.style = doc.styles['MetricValue']
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        r = p.add_run(values[i])
        r.font.size = Pt(size)
        lab = cell.add_paragraph(style=doc.styles['MetricLabel'])
        lab.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        letterspace(lab.add_run(head_texts[i]), 24)
    cant_split(t.rows[0])
    spacer(doc, 6)


def render_table(doc, node):
    head = node.find_all('th')
    body_rows = [tr.find_all('td', recursive=False) for tr in node.find_all('tr')]
    body_rows = [r for r in body_rows if r]
    ncols = max(len(head), max((len(r) for r in body_rows), default=0))
    if not ncols:
        return
    text_rows = [[text_of(c) for c in head]] + [[text_of(c) for c in r] for r in body_rows]
    widths = col_widths(text_rows)
    size = TABLE_PT if ncols <= 4 else TABLE_PT_TIGHT
    right_cols = numeric_columns(text_rows)
    summary_at = next((i for i, r in enumerate(body_rows)
                       if is_summary_row(r)), None)

    # a totals row gets a headline strip of its own above the detail
    if (summary_at is not None and ncols >= 3 and len(body_rows) >= 3
            and 'total' in text_rows[summary_at + 1][0].lower()):
        render_summary_band(doc, text_rows[0], body_rows[summary_at], widths)

    t = doc.add_table(rows=0, cols=ncols)
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    t.autofit = False
    set_tbl_w(t)
    for i, w in enumerate(widths):
        t.columns[i].width = w

    def fill(cells_html, is_head, zebra, summary=False):
        row = t.add_row()
        cant_split(row)
        for i in range(ncols):
            cell = row.cells[i]
            cell.width = widths[i]
            cell_margins(cell, 75 if is_head else 62, 100, 75 if is_head else 62, 100)
            if is_head:
                cell_shade(cell, NAVY)
                cell_borders(cell, top=('single', 4, NAVY), bottom=('single', 4, NAVY),
                             left=('single', 4, NAVY), right=('single', 4, '31425F'))
            elif summary:
                cell_shade(cell, NAVY_SOFT)
                cell_borders(cell, top=('single', 10, NAVY), bottom=('single', 6, NAVY),
                             left=('single', 4, NAVY_SOFT), right=('single', 4, NAVY_SOFT))
            else:
                if zebra:
                    cell_shade(cell, SOFT)
                cell_borders(cell, top=('single', 4, RULE_LT), bottom=('single', 4, RULE_LT),
                             left=('single', 4, RULE_LT), right=('single', 4, RULE_LT))
            p = cell.paragraphs[0]
            p.style = doc.styles['CellHead' if is_head else 'CellText']
            if i in right_cols:
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            if i < len(cells_html):
                lines = tokenize(cells_html[i])
                for j, ln in enumerate(lines):
                    if j:
                        p = cell.add_paragraph(style=doc.styles['CellText'])
                        if i in right_cols:
                            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    emit(p, ln, size=size + (0.5 if summary else 0),
                         color=WHITE if is_head else (NAVY if summary else None),
                         code_size=size - 0.5, bold_all=is_head or summary)
        return row

    if head:
        repeat_header(fill(head, True, False))
    for n, r in enumerate(body_rows):
        fill(r, False, n % 2 == 1, summary=(n == summary_at))
    spacer(doc, 10)


def render_step_table(doc, rows, head=('#', 'Step', 'What happens'),
                      widths=(0.075, 0.30, 0.625)):
    """A plain step table: marker · step name (+ tags) · description.
    rows: [(marker, title, tags[list], detail)] — marker '' renders as a rule-only
    decision row."""
    t = doc.add_table(rows=0, cols=3)
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    t.autofit = False
    set_tbl_w(t)
    cols = [Emu(int(CONTENT_W.emu * f)) for f in widths]
    cols[-1] = Emu(CONTENT_W.emu - cols[0].emu - cols[1].emu)
    for i, w in enumerate(cols):
        t.columns[i].width = w

    hr = t.add_row()
    cant_split(hr)
    for i, label in enumerate(head):
        cell = hr.cells[i]
        cell.width = cols[i]
        cell_shade(cell, NAVY)
        cell_margins(cell, 75, 100, 75, 100)
        cell_borders(cell, top=('single', 4, NAVY), bottom=('single', 4, NAVY),
                     left=('single', 4, NAVY), right=('single', 4, '31425F'))
        para = cell.paragraphs[0]
        para.style = doc.styles['CellHead']
        para.add_run(label)
    repeat_header(hr)

    for n, (marker, title, tags, detail) in enumerate(rows):
        row = t.add_row()
        cant_split(row)
        decision = not str(marker).isdigit()
        for i in range(3):
            cell = row.cells[i]
            cell.width = cols[i]
            cell_margins(cell, 66, 100, 66, 100)
            if decision:
                cell_shade(cell, CALLOUT_BG)
            elif n % 2 == 1:
                cell_shade(cell, SOFT)
            cell_borders(cell, top=('single', 4, RULE_LT), bottom=('single', 4, RULE_LT),
                         left=('single', 4, RULE_LT), right=('single', 4, RULE_LT))
        # marker
        para = row.cells[0].paragraphs[0]
        para.style = doc.styles['CellText']
        r = para.add_run(str(marker) if not decision else '◇')
        r.font.bold = True
        r.font.color.rgb = RGBColor.from_string(COPPER if decision else NAVY)
        # title + inline tags
        para = row.cells[1].paragraphs[0]
        para.style = doc.styles['CellText']
        r = para.add_run(title)
        r.font.bold = True
        r.font.color.rgb = RGBColor.from_string(COPPER if decision else NAVY)
        for tag in tags or ():
            para.add_run('  ')
            r = para.add_run(f' {tag} ')
            r.font.bold = True
            r.font.size = Pt(TABLE_PT - 1)
            r.font.color.rgb = RGBColor.from_string(COPPER)
            run_shade(r, 'F3E6D6')
        # detail
        para = row.cells[2].paragraphs[0]
        para.style = doc.styles['CellText']
        para.add_run(detail)
    spacer(doc, 10)
    return t


def dispatch_flow_rows():
    """The Use Case A assignment loop as table rows, from the live fragment."""
    from opsflow import parse_dispatch_flow
    rows = []
    for item in parse_dispatch_flow()['items']:
        if item['kind'] == 'step':
            rows.append((item['n'], item['title'], item['tags'], item['detail']))
        else:
            arms = item['arms'] + ['', '']
            outcomes = []
            for arm, end in zip(arms, item['ends']):
                sub = f" ({end['sub']})" if end['sub'] else ''
                outcomes.append(f"{arm} → {end['text']}{sub}")
            rows.append(('◇', item['q'], [], '   ·   '.join(outcomes)))
    return rows


def render_dispatch_loop_section(doc, fig_block, timeout_note):
    """Re-typeset §5 A: the figure, then the flow as a step table, then the rule.
    The html export flattens this flow into one run-on paragraph; this rebuilds it
    from the structured source instead."""
    if fig_block is not None:
        render_figure_block(doc, fig_block)
    p = doc.add_paragraph(style=doc.styles['Chip'])
    letterspace(p.add_run('Assignment sequence · trigger to outcome'), 26)
    render_step_table(doc, dispatch_flow_rows())
    if timeout_note:
        callout_text(doc, timeout_note)


def callout_text(doc, text):
    """A one-paragraph copper callout, with a bold "Label:" lead-in if present."""
    t, cell = one_cell(doc, CALLOUT_BG, {
        'left': ('single', 18, COPPER), 'top': ('single', 4, COPPER_LT),
        'bottom': ('single', 4, COPPER_LT), 'right': ('single', 4, COPPER_LT)},
        margins=(110, 150, 110, 140))
    para = cell.paragraphs[0]
    para.style = doc.styles['Callout']
    m = LABEL_LEAD.match(text)
    if m:
        r = para.add_run(m.group(0))
        r.bold = True
        r.font.color.rgb = RGBColor.from_string(NAVY_MID)
        text = text[m.end():]
    para.add_run(text)
    cant_split(t.rows[0])
    spacer(doc, 9)
    return t


def render_metric_strip(doc, pairs):
    """pairs: list of (value, label) rendered as a hairline-ruled strip."""
    n = len(pairs)
    cols = min(n, 4)
    rows = (n + cols - 1) // cols
    col_w = [Emu(int(CONTENT_W.emu / cols))] * cols
    t = doc.add_table(rows=rows, cols=cols)
    t.autofit = False
    set_tbl_w(t)
    for i, c in enumerate(t.columns):
        c.width = col_w[i]
    table_borders(t, inside_h=None, inside_v=None,
                  outer=None)
    for idx, (value, label) in enumerate(pairs):
        cell = t.cell(idx // cols, idx % cols)
        cell.width = col_w[idx % cols]
        cell_margins(cell, 90, 0, 90, 140)
        cell_borders(cell, top=('single', 8, NAVY), bottom=('single', 4, RULE),
                     left=None, right=None)
        p = cell.paragraphs[0]
        p.style = doc.styles['MetricValue']
        p.add_run(value)
        p2 = cell.add_paragraph(style=doc.styles['MetricLabel'])
        letterspace(p2.add_run(label), 28)
    for r in t.rows:
        cant_split(r)
    spacer(doc, 10)


def render_metric_rows(doc, triples):
    """triples: (number, label, detail) — the 'at a glance' count rows."""
    t = doc.add_table(rows=len(triples), cols=3)
    t.autofit = False
    set_tbl_w(t)
    widths = [Emu(int(CONTENT_W.emu * f)) for f in (0.09, 0.28, 0.63)]
    widths[-1] = Emu(CONTENT_W.emu - widths[0].emu - widths[1].emu)
    for i, wd in enumerate(widths):
        t.columns[i].width = wd
    table_borders(t, inside_h=('single', 4, RULE_LT), inside_v=None, outer=None)
    for r_i, (num, label, detail) in enumerate(triples):
        row = t.rows[r_i]
        cant_split(row)
        for c_i in range(3):
            cell = row.cells[c_i]
            cell.width = widths[c_i]
            cell_margins(cell, 70, 0 if c_i == 0 else 90, 70, 90)
            cell_borders(cell, top=('single', 4, RULE_LT) if r_i else ('single', 8, NAVY),
                         bottom=('single', 4, RULE_LT), left=None, right=None)
        p = row.cells[0].paragraphs[0]
        p.style = doc.styles['MetricValue']
        r = p.add_run(num)
        r.font.size = Pt(13)
        r.font.color.rgb = RGBColor.from_string(COPPER)
        p = row.cells[1].paragraphs[0]
        p.style = doc.styles['MetricLabel']
        p.paragraph_format.space_before = Pt(2)
        letterspace(p.add_run(label), 26)
        p = row.cells[2].paragraphs[0]
        p.style = doc.styles['MetricNote']
        p.add_run(detail)
    spacer(doc, 10)


def render_flow(doc, text):
    p = doc.add_paragraph(style=doc.styles['Flow'])
    para_shade(p, FLOW_BG)
    para_borders(p, left=('single', 12, COPPER))
    r = p.add_run(text)
    letterspace(r, 4)
    return p


def render_chip(doc, text):
    p = doc.add_paragraph(style=doc.styles['Chip'])
    m = re.match(r'^([ABC])\s+(.*)$', text)
    if m and m.group(1) in BADGE:
        badge(p, m.group(1), size=8)
        p.add_run('  ')
        text = m.group(2)
    r = p.add_run(text)
    letterspace(r, 26)
    return p


LABEL_LEAD = re.compile(r"^([A-Z][A-Za-z'&/ -]{2,28}):\s")
STEPLINE = re.compile(r'^(\d{1,2})\s+(?=[A-Z↻])')
DASH_RE = re.compile(r'^[-–•]\s+(.*)$')
NUMISH = re.compile(r'^[~<>≤≥]?\d[\d,.\-–—+%~/ ]*\s*[A-Za-z%$]{0,10}\+?$')
BARE_NUM = re.compile(r'^\d{1,2}$')


def classify(txt):
    if DASH_RE.match(txt):
        return 'dash'
    if STEPLINE.match(txt):
        return 'stepline'
    if len(txt) <= 70 and '·' in txt and not txt.endswith(('.', ':')):
        return 'chip'
    if (any(a in txt for a in ARROWS) and len(txt) <= 115
            and not txt.endswith('.') and 2 < txt.count(' ') <= 26):
        return 'flow'
    if len(txt) <= 34 and not txt.endswith(('.', ':')):
        return 'short'
    return 'plain'


def parse_metric_run(lines):
    """Try (num,label,detail) triples then (value,label) pairs. -> (kind, data) or None"""
    txts = [line_text(l) for l in lines]
    # triples: bare number, short label, longer detail
    if len(txts) >= 6 and len(txts) % 3 == 0:
        ok = all(re.fullmatch(r'\d{1,3}', txts[i]) for i in range(0, len(txts), 3))
        if ok:
            return 'triples', [(txts[i], txts[i + 1], txts[i + 2])
                               for i in range(0, len(txts), 3)]
    if len(txts) >= 4 and len(txts) % 2 == 0:
        ok = all(NUMISH.match(txts[i]) for i in range(0, len(txts), 2))
        ok = ok and all(len(txts[i + 1]) <= 26 for i in range(0, len(txts), 2))
        if ok:
            return 'pairs', [(txts[i], txts[i + 1]) for i in range(0, len(txts), 2)]
    return None


DANGLING = {
    'the', 'a', 'an', 'of', 'and', 'or', 'to', 'in', 'on', 'with', 'for', 'by',
    'from', 'is', 'are', 'was', 'were', 'that', 'its', 'their', 'this', 'these',
    'at', 'as', 'into', 'via', 'per', 'plus', 'about', 'across', 'against',
    'between', 'both', 'each', 'every', 'no', 'not', 'than', 'then', 'through',
}


def merge_wrapped(lines):
    """The html export soft-wrapped some sentences; stitch them back together."""
    out = []
    for ln in lines:
        txt = line_text(ln)
        if out:
            prev = line_text(out[-1])
            dangling = prev.rsplit(' ', 1)[-1].lower() in DANGLING
            cont = (txt[:1].islower() or txt[:1] in '(“"' or dangling)
            unfinished = not prev.endswith(('.', ':', '!', '?', ';', '·'))
            if cont and unfinished and classify(prev) == 'plain':
                out[-1] = list(out[-1]) + [Tok(' ', False, False, False, None)] + list(ln)
                continue
        out.append(list(ln))
    return out


def render_annotation_run(doc, texts):
    """Short diagram-branch labels, compressed onto one muted line."""
    p = doc.add_paragraph(style=doc.styles['Chip'])
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(7)
    for i, t in enumerate(texts):
        if i:
            r = p.add_run('  ·  ')
            r.font.color.rgb = RGBColor.from_string(COPPER)
        r = p.add_run(t)
        letterspace(r, 18)
    return p


def render_lines(doc, lines, bullet_level=None):
    """Render the token-lines of one <p>/<li>, re-typesetting flattened structure."""
    lines = merge_wrapped(lines)
    kinds = [classify(line_text(l)) for l in lines]
    i = 0
    first = True
    while i < len(lines):
        txt = line_text(lines[i])
        kind = kinds[i]

        # a run of short lines -> metric strip / count rows / annotation line
        if kind == 'short':
            j = i
            while j < len(lines) and (kinds[j] == 'short' or
                                      (kinds[j] == 'plain' and len(line_text(lines[j])) <= 60)):
                j += 1
            run = lines[i:j]
            parsed = parse_metric_run(run)
            # allow a leading stray label such as "abstract"
            if not parsed and len(run) > 1:
                parsed = parse_metric_run(run[1:])
                if parsed:
                    run = run[1:]
                    i += 1
            if parsed:
                kindname, data = parsed
                if kindname == 'triples':
                    render_metric_rows(doc, data)
                else:
                    render_metric_strip(doc, data)
                i = j
                first = False
                continue
            if len(run) >= 2 and bullet_level is None:
                render_annotation_run(doc, [line_text(l) for l in run])
                i = j
                first = False
                continue

        if kind == 'dash':
            m = DASH_RE.match(txt)
            p = doc.add_paragraph(style=doc.styles['Bullet2' if bullet_level else 'Bullet'])
            r = p.add_run('▪  ' if bullet_level else '●  ')
            r.font.color.rgb = RGBColor.from_string(MUTED if bullet_level else COPPER)
            r.font.size = Pt(7 if bullet_level else 7.5)
            toks = lines[i]
            # strip the leading dash from the first token
            toks = [Tok(re.sub(r'^[-–•]\s+', '', toks[0].text), toks[0].bold, toks[0].italic,
                        toks[0].code, toks[0].href)] + list(toks[1:])
            emit(p, toks)
        elif kind == 'stepline':
            m = STEPLINE.match(txt)
            rest_len = len(txt) - m.end()
            style = 'StepHead' if rest_len <= 64 else 'Normal'
            p = doc.add_paragraph(style=doc.styles[style])
            if style == 'Normal':
                p.paragraph_format.left_indent = Inches(0.32)
                p.paragraph_format.first_line_indent = Inches(-0.32)
                p.paragraph_format.space_before = Pt(5)
            r = p.add_run(m.group(1).rjust(2))
            r.font.name = SANS
            r.font.bold = True
            r.font.color.rgb = RGBColor.from_string(COPPER)
            p.add_run('   ')
            toks = list(lines[i])
            toks[0] = Tok(re.sub(r'^\s*\d{1,2}\s+', '', toks[0].text), toks[0].bold,
                          toks[0].italic, toks[0].code, toks[0].href)
            emit(p, toks)
        elif kind == 'chip':
            render_chip(doc, txt)
        elif kind == 'flow':
            render_flow(doc, txt)
        else:
            if bullet_level is not None and first:
                p = doc.add_paragraph(style=doc.styles['Bullet2' if bullet_level > 0 else 'Bullet'])
                r = p.add_run('▪  ' if bullet_level > 0 else '●  ')
                r.font.color.rgb = RGBColor.from_string(MUTED if bullet_level > 0 else COPPER)
                r.font.size = Pt(7 if bullet_level > 0 else 7.5)
            else:
                p = doc.add_paragraph()
            raw = ''.join(t.text for t in lines[i])
            m = LABEL_LEAD.match(raw)
            if m and not lines[i][0].bold and len(raw) > len(m.group(0)) + 20:
                lead = m.group(0)
                r = p.add_run(lead)
                r.bold = True
                r.font.color.rgb = RGBColor.from_string(NAVY_MID)
                rest = list(lines[i])
                consumed = 0
                new = []
                for t in rest:
                    if consumed >= len(lead):
                        new.append(t)
                        continue
                    take = min(len(t.text), len(lead) - consumed)
                    consumed += take
                    remainder = t.text[take:]
                    if remainder:
                        new.append(Tok(remainder, t.bold, t.italic, t.code, t.href))
                emit(p, new)
            else:
                emit(p, lines[i])
        first = False
        i += 1


def render_list(doc, ul, level=0):
    for li in ul.find_all('li', recursive=False):
        nested = li.find_all(['ul', 'ol'], recursive=False)
        for n in nested:
            n.extract()
        render_lines(doc, tokenize(li), bullet_level=level)
        for n in nested:
            render_list(doc, n, level + 1)


H_MAP = {'h2': 'Heading 1', 'h3': 'Heading 2', 'h4': 'Heading 3', 'h5': 'Heading 4'}
NUM_RE = re.compile(r'^(\d+)\.\s+(.*)$')
ABC_RE = re.compile(r'^((?:[ABC]\s+)+)([A-Z].*)$')


def badge(p, letter, size=None):
    r = p.add_run(f' {letter} ')
    r.font.bold = True
    r.font.small_caps = False
    r.font.color.rgb = RGBColor.from_string(WHITE)
    if size:
        r.font.size = Pt(size)
    run_shade(r, BADGE[letter])
    return r


def render_heading(doc, node, toc, step_num=None):
    style = H_MAP[node.name]
    txt = text_of(node)
    p = doc.add_paragraph(style=doc.styles[style])
    if node.name == 'h2':
        m = NUM_RE.match(txt)
        if m:
            r = p.add_run(m.group(1).zfill(2))
            r.font.color.rgb = RGBColor.from_string(COPPER)
            p.add_run('   ' + m.group(2))
            toc.append((1, m.group(1).zfill(2) + '  ' + m.group(2),
                        f'{len(toc)}|' + m.group(1).zfill(2) + ' ' + m.group(2)))
        else:
            p.add_run(txt)
            toc.append((1, txt, f'{len(toc)}|{txt}'))
        para_borders(p, bottom=('single', 8, NAVY))
        p.paragraph_format.space_after = Pt(14)
        return p

    if step_num:
        r = p.add_run(step_num)
        r.font.bold = True
        r.font.color.rgb = RGBColor.from_string(COPPER)
        p.add_run('   ')

    label = txt
    m = ABC_RE.match(txt) if node.name in ('h3', 'h4') and len(txt) < 90 else None
    if m:
        letters = m.group(1).split()
        for k, letter in enumerate(letters):
            if k:
                p.add_run(' ')
            badge(p, letter)
        p.add_run('  ')
        label = m.group(2)
        p.add_run(label)
    else:
        for ln in tokenize(node):
            emit(p, ln)
            break
    if node.name == 'h3':
        para_borders(p, bottom=('single', 4, RULE))
        p.paragraph_format.space_after = Pt(7)
        toc.append((2, label, f'{len(toc)}|{label}'))
    return p


# ------------------------------------------------------------- page furniture
def find_logo():
    for pat in ('logo.png', 'logo.jpg', 'logo*.png', 'logo*.jpg', 'logo*.jpeg',
                '*logo*.png', 'techjays*.png', 'assets/logo*.png', 'assets/*logo*.png',
                'assets/*logo*.jpg'):
        hits = sorted(ROOT.glob(pat))
        if hits:
            return hits[0]
    return None


LOGO = find_logo()


def wordmark(paragraph, size=10.5, color=NAVY, accent=COPPER, space=24):
    r = paragraph.add_run('tech')
    r.font.name, r.font.size, r.font.bold = SANS, Pt(size), True
    r.font.color.rgb = RGBColor.from_string(color)
    letterspace(r, space)
    r = paragraph.add_run('jays')
    r.font.name, r.font.size, r.font.bold = SANS, Pt(size), True
    r.font.color.rgb = RGBColor.from_string(accent)
    letterspace(r, space)


def two_col_strip(container, styles, left_build, right_build, rule='bottom',
                  width=CONTENT_W):
    """A borderless 2-column table used for header / footer rows (tab-stop free)."""
    t = container.add_table(rows=1, cols=2, width=width)
    t.autofit = False
    set_tbl_w(t, width)
    half = Emu(int(width.emu / 2))
    for c in t.columns:
        c.width = half
    edge = {'top': None, 'bottom': None, 'left': None, 'right': None}
    if rule == 'bottom':
        edge['bottom'] = ('single', 4, RULE)
    elif rule == 'top':
        edge['top'] = ('single', 4, RULE)
    for idx, build in enumerate((left_build, right_build)):
        cell = t.cell(0, idx)
        cell.width = half
        cell_borders(cell, **edge)
        cell_margins(cell, 20, 0, 40, 0)
        p = cell.paragraphs[0]
        p.style = styles['HeadFoot']
        if idx:
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        build(p)
    return t


def build_header(section, styles, width=CONTENT_W):
    """The brand banner: techjays mark plus the gradient rule, full text width."""
    hdr = section.header
    hdr.is_linked_to_previous = False
    p = hdr.paragraphs[0]
    p.text = ''
    p.style = styles['HeadFoot']
    p.paragraph_format.space_after = Pt(2)
    banner = BRAND / 'header-banner.png'
    if banner.exists():
        p.add_run().add_picture(str(banner), width=width)
    else:
        wordmark(p, size=9, space=20)
    return p


def build_footer(section, styles, width=CONTENT_W):
    ftr = section.footer
    ftr.is_linked_to_previous = False
    ftr.paragraphs[0].text = ''

    def left(p):
        r = p.add_run('Confidential  —  ClimatePros / techjays')
        r.font.size, r.font.name = Pt(8), SANS
        r.font.color.rgb = RGBColor.from_string(MUTED)

    def right(p):
        def g(txt, bold=False, color=MUTED):
            r = p.add_run(txt)
            r.font.size, r.font.name, r.font.bold = Pt(8), SANS, bold
            r.font.color.rgb = RGBColor.from_string(color)
        g('Page ')
        for r in add_field(p, 'PAGE'):
            r.font.size, r.font.name, r.font.bold = Pt(8), SANS, True
            r.font.color.rgb = RGBColor.from_string(NAVY)
        g(' of ')
        for r in add_field(p, 'NUMPAGES'):
            r.font.size, r.font.name = Pt(8), SANS
            r.font.color.rgb = RGBColor.from_string(MUTED)

    two_col_strip(ftr, styles, left, right, rule='top', width=width)
    ftr.paragraphs[-1].paragraph_format.space_after = Pt(0)
    for r in ftr.paragraphs[-1].runs:
        r.font.size = Pt(2)


def page_setup(section, landscape=False):
    if landscape:
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width = Inches(11)
        section.page_height = Inches(8.5)
        section.left_margin = Inches(0.7)
        section.right_margin = Inches(0.7)
        section.top_margin = Inches(0.7)
        section.bottom_margin = Inches(0.65)
    else:
        section.orientation = WD_ORIENT.PORTRAIT
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)
        section.top_margin = Inches(0.95)
        section.bottom_margin = Inches(0.9)
    section.header_distance = Inches(0.45)
    section.footer_distance = Inches(0.4)
    return Emu(section.page_width.emu - section.left_margin.emu - section.right_margin.emu)


def new_page_section(doc, landscape=False, center=False):
    """Start a fresh page (optionally rotated) and re-hang the header/footer."""
    sec = doc.add_section(WD_SECTION.NEW_PAGE)
    width = page_setup(sec, landscape=landscape)
    # a new section inherits the previous sectPr, including its page-number
    # restart — drop it so numbering runs continuously through the document
    stale = sec._sectPr.find(qn('w:pgNumType'))
    if stale is not None:
        sec._sectPr.remove(stale)
    build_header(sec, doc.styles, width=width)
    build_footer(sec, doc.styles, width=width)
    return sec, width


def restart_numbering(section, start=1):
    put(section._sectPr, _el('w:pgNumType', start=start))


# ------------------------------------------------------------------- cover page
def full_page_background(doc, image, page_w=Inches(8.5), page_h=Inches(11)):
    """Anchor artwork behind the text at page origin, so the cover can be typeset
    on top of it. python-docx only makes inline shapes, so the inline run is
    rewritten into a floating anchor with behindDoc set."""
    para = doc.add_paragraph()
    para.paragraph_format.space_after = Pt(0)
    para.paragraph_format.line_spacing = Pt(1)
    run = para.add_run()
    run.font.size = Pt(1)
    run.add_picture(str(image), width=page_w, height=page_h)
    inline = para._p.find('.//' + qn('wp:inline'))
    if inline is None:
        return para
    anchor = OxmlElement('wp:anchor')
    for k, v in (('distT', '0'), ('distB', '0'), ('distL', '0'), ('distR', '0'),
                 ('simplePos', '0'), ('relativeHeight', '0'), ('behindDoc', '1'),
                 ('locked', '0'), ('layoutInCell', '1'), ('allowOverlap', '1')):
        anchor.set(k, v)
    simple = OxmlElement('wp:simplePos'); simple.set('x', '0'); simple.set('y', '0')
    anchor.append(simple)
    for tag, rel, off in (('wp:positionH', 'page', 0), ('wp:positionV', 'page', 0)):
        pos = OxmlElement(tag)
        pos.set('relativeFrom', rel)
        o = OxmlElement('wp:posOffset')
        o.text = str(off)
        pos.append(o)
        anchor.append(pos)
    for child in list(inline):
        anchor.append(child)
    wrap = OxmlElement('wp:wrapNone')
    anchor.insert(list(anchor).index(anchor.find(qn('wp:extent'))) + 1, wrap) \
        if anchor.find(qn('wp:extent')) is not None else anchor.append(wrap)
    inline.getparent().replace(inline, anchor)
    return para


def build_cover(doc, meta, usecases, version):
    art = BRAND / 'cover-art.jpg'
    if art.exists():
        full_page_background(doc, art)

    # the title block sits over the artwork's pale upper field
    vgap(doc, 112)
    p = doc.add_paragraph(style=doc.styles['CoverTitle'])
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run('ClimatePros')
    p = doc.add_paragraph(style=doc.styles['CoverSub'])
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('Unified Architecture')
    r.font.name = SERIF
    r.font.size = Pt(17)
    r.font.bold = True
    r.font.color.rgb = RGBColor.from_string(NAVY_MID)
    p.paragraph_format.space_after = Pt(16)

    p = doc.add_paragraph(style=doc.styles['CoverSub'])
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for i, line in enumerate(usecases):
        m = re.match(r'^Use Case ([ABC])\s*—\s*(.*)$', line)
        if i:
            sep = p.add_run('     ')
            sep.font.size = Pt(10)
        if m:
            r = p.add_run(m.group(1) + '  ')
            r.font.bold = True
            r.font.size = Pt(10)
            r.font.color.rgb = RGBColor.from_string(VIOLET)
            r = p.add_run(m.group(2))
        else:
            r = p.add_run(line)
        r.font.size = Pt(10)
        r.font.color.rgb = RGBColor.from_string(NAVY_MID)

    p = doc.add_paragraph(style=doc.styles['CoverSub'])
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(10)
    r = p.add_run(version)
    r.font.size = Pt(9.5)
    r.font.bold = True
    r.font.color.rgb = RGBColor.from_string(VIOLET_DEEP)

    # white-on-navy strip sits inside the artwork's lower field
    vgap(doc, 300)
    for label, value in list(meta.items())[:4]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.line_spacing = 1.0
        r = p.add_run(f'{label.upper()}   ')
        r.font.name = SANS
        r.font.size = Pt(7.5)
        r.font.bold = True
        r.font.color.rgb = RGBColor.from_string('9BB0D6')
        letterspace(r, 30)
        r = p.add_run(value)
        r.font.name = SANS
        r.font.size = Pt(9)
        r.font.color.rgb = RGBColor.from_string(WHITE)


# ------------------------------------------------------------------ front matter
def render_front(doc, blocks):
    for node in blocks:
        if node.name == 'p':
            txt = text_of(node)
            if txt in ('Document Control', 'How to Read This Document', 'Approval'):
                p = doc.add_paragraph(style=doc.styles['BlockTitle'])
                p.add_run(txt)
                para_borders(p, bottom=('single', 8, NAVY))
                continue
            if txt.startswith('Approved for'):
                p = doc.add_paragraph()
                p.paragraph_format.space_after = Pt(15)
                p.add_run(txt).font.size = Pt(10)
                continue
            render_lines(doc, tokenize(node))
        elif node.name == 'table':
            render_table(doc, node)
        elif node.name in ('ul', 'ol'):
            render_list(doc, node)
        elif node.name == 'pre':
            render_code(doc, node)


def render_toc(doc, entries, page_map):
    p = doc.add_paragraph(style=doc.styles['BlockTitle'])
    p.paragraph_format.page_break_before = True
    p.paragraph_format.space_before = Pt(0)
    p.add_run('Contents')
    para_borders(p, bottom=('single', 8, NAVY))
    p.paragraph_format.space_after = Pt(14)
    for level, label, key in entries:
        style = 'TocH1' if level == 1 else 'TocH2'
        tp = doc.add_paragraph(style=doc.styles[style])
        tp.paragraph_format.tab_stops.add_tab_stop(
            CONTENT_W, WD_TAB_ALIGNMENT.RIGHT)
        tp.add_run(label)
        page = page_map.get(key, '')
        r = tp.add_run('\t' + (str(page) if page else '00'))
        r.font.color.rgb = RGBColor.from_string(MUTED if level == 2 else NAVY)
        r.font.bold = (level == 1)


DISPATCH_H3 = 'Auto-Dispatch — Assignment Loop'
SNS_H4 = 'Amazon SNS round-trip'


def rebuild_dispatch_subsection(doc, blocks, idx, toc):
    """§5 A comes out of the html export as one run-on paragraph plus five
    heading/paragraph pairs. Rebuild it from sections/05.html and report how many
    source blocks were consumed."""
    end = next((j for j in range(idx + 1, len(blocks))
                if blocks[j].name == 'h3'), len(blocks))
    region = blocks[idx:end]

    render_heading(doc, region[0], toc)
    intro = next((b for b in region[1:] if b.name == 'p'), None)
    if intro is not None:
        render_lines(doc, tokenize(intro))

    fig = next((b for b in region if b.name == 'blockquote'), None)
    # the timeout guardrail is the tail line of the flattened flow paragraph
    timeout = ''
    for b in region:
        if b.name == 'p':
            for line in tokenize(b):
                txt = line_text(line)
                if txt.startswith('Timeout rule'):
                    timeout = txt
    render_dispatch_loop_section(doc, fig, timeout)

    sns_head = next((b for b in region if b.name == 'h4'
                     and SNS_H4 in text_of(b)), None)
    if sns_head is not None:
        render_heading(doc, sns_head, toc)
        lead = None
        for j, b in enumerate(region):
            if b is sns_head:
                lead = region[j + 1] if j + 1 < len(region) else None
                break
        offer_n = next((r[0] for r in dispatch_flow_rows()
                        if str(r[1]).lower().startswith('offer to the top')), None)
        if lead is not None and lead.name == 'p':
            lines = [ln for ln in tokenize(lead)
                     if not BARE_NUM.fullmatch(line_text(ln))]
            if lines:
                txt = ' '.join(line_text(ln) for ln in lines)
                if offer_n:      # the export points at the old inner numbering
                    txt = re.sub(r"step \d+'s push", f"step {offer_n}'s push", txt)
                para = doc.add_paragraph()
                para.add_run(txt)
        rows, i = [], 0
        for j in range(len(region)):
            b = region[j]
            if b.name == 'h4' and b is not sns_head:
                body = region[j + 1] if j + 1 < len(region) else None
                detail = text_of(body) if body is not None and body.name == 'p' else ''
                detail = re.sub(r'\s+\d+$', '', detail)   # stray trailing step number
                i += 1
                rows.append((i, text_of(b), [], detail))
        render_step_table(doc, rows)
    return end - idx


def render_blocks(doc, blocks, toc, section_no=None, in_override=False):
    pending_num = None
    skip_to = 0
    for idx, node in enumerate(blocks):
        if idx < skip_to:
            continue
        n = node.name
        if n == 'h3' and DISPATCH_H3 in text_of(node):
            skip_to = idx + rebuild_dispatch_subsection(doc, blocks, idx, toc)
            continue
        if n == 'h3' and not in_override:
            md = sub_override(section_no, text_of(node))
            if md is not None:
                render_blocks(doc, markdown_blocks(md, MD_SUB_HEADING_MAP), toc,
                              section_no, in_override=True)
                skip_to = next((j for j in range(idx + 1, len(blocks))
                                if blocks[j].name in ('h2', 'h3')), len(blocks))
                continue
        if n == 'p':
            img = node.find('img')
            if img is not None and len(node.find_all(True)) == 1:
                if render_md_figure(doc, img):
                    continue
        if n in H_MAP:
            render_heading(doc, node, toc, step_num=pending_num)
            pending_num = None
        elif n == 'p':
            if not text_of(node):
                continue
            lines = tokenize(node)
            # a trailing bare number belongs to the step heading that follows
            nxt = blocks[idx + 1].name if idx + 1 < len(blocks) else None
            if lines and nxt in ('h4', 'h5') and BARE_NUM.fullmatch(line_text(lines[-1])):
                pending_num = line_text(lines[-1])
                lines = lines[:-1]
            if not lines:
                continue
            render_lines(doc, lines)
        elif n == 'table':
            render_table(doc, node)
        elif n in ('ul', 'ol'):
            render_list(doc, node)
        elif n == 'pre':
            render_code(doc, node)
        elif n == 'blockquote':
            strongs = node.find_all('strong')
            is_fig = any(FIG_RE.match(text_of(s.parent)) for s in strongs) if strongs else False
            render_figure_block(doc, node) if is_fig else render_callout(doc, node)
        elif n == 'hr':
            p = doc.add_paragraph()
            para_borders(p, bottom=('single', 4, RULE))
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(8)
            p.add_run('').font.size = Pt(2)


# ------------------------------------------------------- markdown-backed sections
MD_DIR = ROOT / 'md-files'
MD_SUB_DIR = MD_DIR / 'subsections'
MD_HEADING_MAP = {'h1': 'h2', 'h2': 'h3', 'h3': 'h4', 'h4': 'h5', 'h5': 'h5',
                  'h6': 'h5'}
# a subsection file's own h1 is the h3 it replaces, so everything shifts one level
MD_SUB_HEADING_MAP = {'h1': 'h3', 'h2': 'h4', 'h3': 'h5', 'h4': 'h5', 'h5': 'h5',
                      'h6': 'h5'}


def slugify(text):
    return re.sub(r'-+', '-', re.sub(r'[^a-z0-9]+', '-', text.lower())).strip('-')


def sub_override(section_no, heading_text):
    """md-files/subsections/NN-slug.md replaces one h3 subtree of section NN."""
    if not MD_SUB_DIR.is_dir() or section_no is None:
        return None
    target = slugify(heading_text)
    for path in sorted(MD_SUB_DIR.glob(f'{int(section_no):02d}-*.md')):
        stem = path.stem.split('-', 1)[1] if '-' in path.stem else path.stem
        if target.startswith(stem) or stem in target:
            return path
    return None


def md_override(section_no):
    """md-files/NN-anything.md replaces the exported html for that section."""
    if not MD_DIR.is_dir():
        return None
    for pattern in (f'{int(section_no):02d}-*.md', f'{int(section_no)}-*.md'):
        hits = sorted(MD_DIR.glob(pattern))
        if hits:
            return hits[0]
    return None


def markdown_blocks(path, level_map=None):
    """Markdown -> the same block soup the html export produces."""
    try:
        import markdown
    except ImportError:                                        # noqa: BLE001
        sys.exit("markdown is needed for md-files sections: pip install markdown")
    html = markdown.markdown(
        path.read_text(encoding='utf-8'),
        extensions=['tables', 'fenced_code', 'sane_lists', 'attr_list'])
    soup = BeautifulSoup(html, 'html.parser')
    level_map = level_map or MD_HEADING_MAP
    for tag in soup.find_all(list(level_map)):
        tag.name = level_map[tag.name]
    # fenced code arrives as <pre><code>…</code></pre>; the renderer wants <pre>
    for pre in soup.find_all('pre'):
        code = pre.find('code')
        if code is not None:
            code.unwrap()
    return [c for c in soup.children if c.name]


def render_body(doc, blocks, toc):
    """Walk the export section by section, swapping in md-files where present."""
    i = 0
    while i < len(blocks):
        end = next((j for j in range(i + 1, len(blocks))
                    if blocks[j].name == 'h2'), len(blocks))
        node = blocks[i]
        md = None
        if node.name == 'h2':
            m = NUM_RE.match(text_of(node))
            md = md_override(m.group(1)) if m else None
        section_no = None
        if node.name == 'h2':
            mm = NUM_RE.match(text_of(node))
            section_no = mm.group(1) if mm else None
        if md is not None:
            render_blocks(doc, markdown_blocks(md), toc, section_no)
        else:
            render_blocks(doc, blocks[i:end], toc, section_no)
        i = end


# --------------------------------------------------------------------- assembly
def parse_source():
    soup = BeautifulSoup(SRC.read_text(encoding='utf-8'), 'html.parser')
    kids = [c for c in soup.body.children if c.name]
    hr_idx = [i for i, c in enumerate(kids) if c.name == 'hr']
    cover_block = kids[:hr_idx[0]]
    front_block = kids[hr_idx[0] + 1:hr_idx[1]]
    after = kids[hr_idx[1] + 1:]
    start = next((i for i, c in enumerate(after) if c.name == 'h2'), 0)
    return cover_block, front_block, after[start:]


def squash_section_breaks(doc):
    """python-docx parks each outgoing section's properties in an empty paragraph.
    At body size that stray line can spill onto a sheet of its own, so pin every
    one of them to a 1pt line."""
    body = doc.element.body
    count = 0
    for para in body.findall(qn('w:p')):
        pPr = para.find(qn('w:pPr'))
        if pPr is None or pPr.find(qn('w:sectPr')) is None:
            continue
        put(pPr, _el('w:spacing', before=0, after=0, line=20, lineRule='exact'))
        rPr = pPr.find(qn('w:rPr'))
        if rPr is None:
            rPr = OxmlElement('w:rPr')
            put(pPr, rPr)
        for tag in ('w:sz', 'w:szCs'):
            old = rPr.find(qn(tag))
            if old is not None:
                rPr.remove(old)
            rPr.append(_el(tag, val=2))
        count += 1
    return count


def embed_display_font(docx_path):
    """Embed the Poppins family so the document renders as designed on any
    machine — same approach the sample uses (plain .ttf parts + fontTable refs)."""
    faces = [('embedRegular', 'Poppins-regular.ttf'), ('embedBold', 'Poppins-bold.ttf'),
             ('embedItalic', 'Poppins-italic.ttf'),
             ('embedBoldItalic', 'Poppins-boldItalic.ttf')]
    faces = [(tag, name) for tag, name in faces if (FONT_DIR / name).exists()]
    if not faces:
        return False

    import zipfile
    tmp = docx_path.with_suffix('.embed.docx')
    with zipfile.ZipFile(docx_path) as zin, \
            zipfile.ZipFile(tmp, 'w', zipfile.ZIP_DEFLATED) as zout:
        names = set(zin.namelist())
        rels_part = 'word/_rels/fontTable.xml.rels'
        for item in zin.infolist():
            if item.filename == rels_part:
                continue                      # rewritten below with the font rels
            data = zin.read(item.filename)
            if item.filename == '[Content_Types].xml':
                text = data.decode('utf-8')
                if 'Extension="ttf"' not in text:
                    text = text.replace('<Types ', '<Types ', 1)
                    text = text.replace('</Types>',
                                        '<Default Extension="ttf" '
                                        'ContentType="application/x-font-ttf"/></Types>')
                data = text.encode('utf-8')
            elif item.filename == 'word/fontTable.xml':
                text = data.decode('utf-8')
                rels = ''.join(
                    f'<w:{tag} w:fontKey="{{00000000-0000-0000-0000-000000000000}}" '
                    f'r:id="rIdFont{i}" w:subsetted="0"/>'
                    for i, (tag, _) in enumerate(faces, start=1))
                entry = f'<w:font w:name="Poppins">{rels}</w:font>'
                if '<w:font w:name="Poppins">' in text:
                    text = re.sub(r'<w:font w:name="Poppins">.*?</w:font>', entry,
                                  text, flags=re.S)
                else:
                    text = text.replace('</w:fonts>', entry + '</w:fonts>')
                data = text.encode('utf-8')
            elif item.filename == 'word/settings.xml':
                text = data.decode('utf-8')
                if 'embedTrueTypeFonts' not in text:
                    # CT_Settings is a sequence: this element precedes proofState
                    for anchor in ('<w:proofState', '<w:defaultTabStop', '</w:settings>'):
                        if anchor in text:
                            text = text.replace(
                                anchor, '<w:embedTrueTypeFonts/>' + anchor, 1)
                            break
                data = text.encode('utf-8')
            zout.writestr(item, data)

        rel_items = ''.join(
            f'<Relationship Id="rIdFont{i}" Type="http://schemas.openxmlformats.org/'
            f'officeDocument/2006/relationships/font" Target="fonts/{name}"/>'
            for i, (_, name) in enumerate(faces, start=1))
        if rels_part in names:
            merged = zin.read(rels_part).decode('utf-8').replace(
                '</Relationships>', rel_items + '</Relationships>')
        else:
            merged = ('<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
                      '<Relationships xmlns="http://schemas.openxmlformats.org/'
                      f'package/2006/relationships">{rel_items}</Relationships>')
        zout.writestr(rels_part, merged)
        for _, name in faces:
            zout.writestr(f'word/fonts/{name}', (FONT_DIR / name).read_bytes())
    tmp.replace(docx_path)
    return True


def build(page_map):
    cover_block, front_block, body_block = parse_source()
    ptexts = [text_of(c) for c in cover_block if c.name == 'p']
    version = ptexts[0]
    usecases = [s.strip() for s in re.split(r'\s+·\s+', ptexts[1]) if s.strip()]
    labels = ['PREPARED FOR', 'PREPARED BY', 'ISSUED', 'COMPANION TO', 'CLASSIFICATION']
    pretty = {'PREPARED FOR': 'Prepared for', 'PREPARED BY': 'Prepared by',
              'ISSUED': 'Issued', 'COMPANION TO': 'Companion to',
              'CLASSIFICATION': 'Classification'}
    meta = {}
    for i, t in enumerate(ptexts):
        if t in labels and i + 1 < len(ptexts):
            meta[pretty[t]] = ptexts[i + 1]

    doc = Document()
    build_styles(doc)
    cover_section = doc.sections[0]
    page_setup(cover_section)
    build_cover(doc, meta, usecases, version)

    body_section = doc.add_section(WD_SECTION.NEW_PAGE)
    page_setup(body_section)
    restart_numbering(body_section, 1)
    build_header(body_section, doc.styles)
    build_footer(body_section, doc.styles)
    cover_section.header.is_linked_to_previous = False
    cover_section.footer.is_linked_to_previous = False

    render_front(doc, front_block)

    # TOC needs the entry list, so render the body into a throwaway pass first
    probe = Document()
    build_styles(probe)
    toc = []
    render_body(probe, body_block, toc)

    render_toc(doc, toc, page_map)
    cover_block, front_block, body_block = parse_source()   # tokenizer mutates nodes
    render_body(doc, body_block, [])

    # no auto-hyphenation: keeps compound technical terms whole.
    # CT_Settings is a sequence, so these have to sit right after defaultTabStop.
    settings = doc.settings.element
    zoom = settings.find(qn('w:zoom'))
    if zoom is not None and zoom.get(qn('w:percent')) is None:
        zoom.set(qn('w:percent'), '100')
    for tag, val in (('w:autoHyphenation', 'false'), ('w:doNotHyphenateCaps', 'true')):
        put(settings, _el(tag, val=val))

    squash_section_breaks(doc)

    core = doc.core_properties
    core.title = DOC_TITLE
    core.subject = 'Stream 1 — Unified Architecture (Use Cases A, B, C)'
    core.author = 'techjays'
    core.category = 'Architecture'
    core.comments = 'Confidential — ClimatePros / techjays'
    return doc, toc


def pdf_page_map(docx_path, entries):
    """Render to PDF and find the page each heading lands on."""
    WORK.mkdir(exist_ok=True)
    for f in WORK.glob('*'):
        f.unlink()
    tmp = WORK / docx_path.name
    shutil.copy(docx_path, tmp)
    subprocess.run(['soffice', '--headless', '--norestore', '--convert-to', 'pdf',
                    '--outdir', str(WORK), str(tmp)],
                   check=True, capture_output=True, timeout=600)
    pdf = tmp.with_suffix('.pdf')
    txt = subprocess.run(['pdftotext', '-layout', str(pdf), '-'],
                         check=True, capture_output=True, text=True).stdout
    pages = txt.split('\f')
    # drop TOC rows (a label, a wide gap, then a page number) and note which
    # pages they live on, so headings are only matched after the contents list
    page_lines, toc_pages = [], set()
    for pno, pg in enumerate(pages, start=1):
        keep, dropped = [], 0
        for ln in pg.split('\n'):
            s = ln.rstrip()
            if re.search(r'\s{3,}\d{1,3}$', s):
                dropped += 1
                continue
            keep.append(re.sub(r'\s+', ' ', s).strip())
        if dropped >= 4 and pno <= 8:      # the contents list sits in the front matter
            toc_pages.add(pno)
        page_lines.append(keep)
    first_body = (max(toc_pages) + 1) if toc_pages else 1
    page_map = {}
    for level, label, key in entries:
        needle = re.sub(r'\s+', ' ', key.split('|', 1)[1]).strip()
        if not needle:
            continue
        for pno in range(first_body, len(pages) + 1):
            if any(is_heading_line(ln, needle) for ln in page_lines[pno - 1]):
                page_map[key] = pno - 1   # cover is unnumbered; body restarts at 1
                break
    return page_map, len([x for x in pages if x.strip()])


def is_heading_line(line, needle):
    """A heading occupies its own line; prose mentioning the same words does not.
    Allow a short prefix for the A/B/C badge letter."""
    ln = line.strip()
    if not ln or needle not in ln:
        return False
    slack = len(ln) - len(needle)
    if slack > 4:
        return False
    return ln.endswith(needle) or ln.startswith(needle)


def main():
    doc, entries = build({})
    doc.save(OUT)
    embed_display_font(OUT)
    try:
        page_map, npages = pdf_page_map(OUT, entries)
        found = sum(1 for _, _, k in entries if k in page_map)
        doc, entries = build(page_map)
        doc.save(OUT)
        embed_display_font(OUT)
        print(f'pass 2 done · TOC page numbers resolved {found}/{len(entries)} · '
              f'{npages} pages')
    except Exception as exc:                                   # noqa: BLE001
        print(f'!! page-number pass skipped: {exc}', file=sys.stderr)
    print(f'wrote {OUT} ({OUT.stat().st_size:,} bytes) logo={LOGO}')


if __name__ == '__main__':
    main()
